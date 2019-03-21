# -*- coding: utf-8 -*-
import tornado.ioloop
import tornado.web
import torndb
import pandas as pd
import datetime
import docx
import os
import undoandtrans

from database import database
from authenticate import *
from function import *
import sys
reload(sys)
sys.setdefaultencoding('utf-8')


global ia
global da
ia=''
da=''



# BaseHandler 基类覆写 get_current_user
# 覆写后 RequestHandler 的current_user成员会有值(稍后解释实现源码)
# 这里简单地判断请求带的 secure cookie 是否带有 speedfairuser属性的值

class BaseHandler(tornado.web.RequestHandler):
    def get_current_user(self):
        return self.get_secure_cookie("speedfairuser")


################################################################################################
#用户登陆页面。提交登陆表单后，后台进行数据验证。验证通过后在在客户端设置COOKIE值。然后重定向至主页面。
class SigninHandler(BaseHandler):
    
    def get(self):
        if not self.current_user:
        	self.render("static/signin.html")  
        else:
		self.render("static/index.html")
               

    def post(self):
        
        username=self.get_argument('name')
        password=self.get_argument('password')
        
        if database:

            u=database.query("SELECT * FROM namepwd WHERE username=%s",username);
            if u:
                
            	pw=u[0]['pwd']
                if pw==password:
			self.set_secure_cookie("speedfairuser", username,expires_days=None)
                  	self.redirect("/")
			#self.render("static/index.html")
                else:self.render("static/msg.html",msg='bad password')
            else:self.render("static/msg.html",msg='no this username')
                
                
                
                
        else:self.render("static/msg.html",msg='some wrong with database')


###############################################################################################################

#########################################################################################

class Setiada(tornado.web.RequestHandler):
	@checkadmin
	def get(self):
		self.render("static/setiada.html")


	def post(self):
		global ia
		global da

			
 		sia=self.get_argument('sia')
 		sda=self.get_argument('sda')
		ia=sia
		da=sda
		self.write("now ia :")
		self.write(ia)
		self.write("now da :")
		self.write(da)

#########################################################################################

#########################################################################################

class MainHandler(tornado.web.RequestHandler):
    @authenticated
    def get(self):

	self.render("static/index.html")



#########################################################################################

class Addnewuser(tornado.web.RequestHandler):
    @checkadmin
    def get(self):
	self.render("static/addnewuser.html")

#########################################################################################

class Addnewuserp(tornado.web.RequestHandler):

    @checkadmin
    def post(self):
        nusername=self.get_argument('nname')
        npassword=self.get_argument('npassword')
        nclassv=self.get_argument('classv')
        sqli="insert into namepwd (username,pwd,level) values (%s,%s,%s)"
        database.insert(sqli,nusername,npassword,nclassv)
        self.render("static/msg.html",msg="it is ok!")
#########################################################################################

class Viewuser(tornado.web.RequestHandler):
	@checkadmin
	def get(self):
        	sql="select * from namepwd"
        	dr=database.query(sql)
		df=pd.DataFrame(dr)
		df=df.reset_index(drop=True)
		ldf=df.index.tolist()
		self.render("static/viewunresult.html",df=df,ldf=ldf)

#########################################################################################
class Changepwd(tornado.web.RequestHandler):

	@authenticated
	def get(self):
		self.render("static/changepwd.html")
	@authenticated
	def post(self):
		name=self.get_secure_cookie('speedfairuser')	
        	newpwd=self.get_argument('newpwd')
        	sql="update namepwd set pwd = %s where username = %s"
        	database.execute(sql,newpwd,name)
		self.render("static/msg.html",msg="密码更改成功.请牢记自己密码.")

#########################################################################################


#########################################################################################
class Deleteuser(tornado.web.RequestHandler):

    @checkadmin
    def get(self):
	self.render("static/deleteuser.html")
    @checkadmin
    def post(self):	
        name=self.get_argument('name')
        sql="delete from namepwd where username='"+name+"'"
        database.execute(sql)
	self.render("static/msg.html",msg="deleting is ok")

#########################################################################################
class Toexcel(tornado.web.RequestHandler):

    @checkadmin
    def get(self):
	und=database.query("select * from stocktable   order by right(exda,4) asc ,left(exda,2) asc")
	pdf=pd.DataFrame(und)
	pdfd=pdf.set_index("id")
	pdfd.to_excel("datasfc.xls")
	self.render("static/index.html")
	

#########################################################################################



#########################################################################################
#####库存查询
class Querystock(tornado.web.RequestHandler):
    @authenticated
    def get(self):	
	
	self.render("static/querystock.html")
    @authenticated
    def post(self):
	kword=self.get_argument('kword')
	waname=self.get_argument('wana')
	opkw=self.get_argument('opkw')
	if opkw=="banu":
		nameq="%"+kword+"%"
		sql="select * from stocktable where banu like %s   order by mana asc, past asc,right(exda,4) asc ,left(exda,2) asc"
		rec=database.query(sql,nameq)
		if len(rec):
			dfda=rectodataframe(rec)
			#dfd=dfda.sort_values(by='exda'.decode('utf-8'),ascending=True)
			dfd=dfda
			if  len(dfd):
		
				if  not waname=="all":
					dfw=dfd[dfd.wana==waname]
				else:dfw=dfd
				df=dfw.reset_index(drop=True)
				quantity=(df.quan*df.past).sum()
				amount=df.quan.sum()
				ldf=df.index.tolist()
				self.render("static/rqstockresult.html",df=df,quantity=quantity,amount=amount,ldf=ldf)

			else:self.render("static/msg.html",msg="您查询的结果为空.")
		else:self.render("static/msg.html",msg="您查询的结果为空.")
	elif opkw=="mana":
		nameq="%"+kword+"%"
		sql="select * from stocktable where mana like %s   order by mana , past asc,right(exda,4) asc ,left(exda,2) asc"
		rec=database.query(sql,nameq)
		if len(rec):
			dfda=rectodataframe(rec)
			#dfd=dfda.sort_values(by='exda'.decode('utf-8'),ascending=True)
			dfd=dfda
			if  len(dfd):
		
				if  not waname=="all":
					dfw=dfd[dfd.wana==waname]
				else:dfw=dfd
	
				df=dfw.reset_index(drop=True)
				quantity=(df.quan*df.past).sum()
				amount=df.quan.sum()
				ldf=df.index.tolist()
				self.render("static/rqstockresult.html",df=df,quantity=quantity,amount=amount,ldf=ldf)
		
			
			else:self.render("static/msg.html",msg="您查询的结果为空.")
		else:self.render("static/msg.html",msg="您查询的结果为空.")
	else:pass
		



#########################################################################################
class Rqfrpr(tornado.web.RequestHandler):
    @authenticated
    def get(self):	
	self.render("static/rqfrpr.html")
    @authenticated
    def post(self):

        startd=self.get_argument('startd')
	stopd=self.get_argument('stopd')
	frpr=self.get_argument('frpr')
	cuna=self.get_argument('cuna')
	#cuna=cuna.decode('utf8')
	rcuna="%"+cuna+"%"
	try:
		startd=startd.split("-")[0]+startd.split("-")[1]+startd.split("-")[2]
		stopd=stopd.split("-")[0]+stopd.split("-")[1]+stopd.split("-")[2]
	except:
		startd=strtoday()[0:5]+"0101"
		stopd=str(int(strtoday()[0:5])+1)+"0101"
	
	if frpr=="fr":
		rdfr=databasetabletodataframe2("freighttable",rcuna)
		if len(rdfr)>0:
			rdfr=rdfr[rdfr.date>startd]
			rdfr=rdfr[rdfr.date<stopd]

			rdfr1=rdfr.reset_index(drop=True)
			df=rdfr1

			l=len(df)

			if l>0:

				amount=round(df.frei.sum(),2)
				ldf=df.index.tolist()
				self.render("static/rqfrresult.html",df=df,amount=amount,ldf=ldf)
			else:
				self.render("static/msg.html",msg="查询结果为空.")
		else:
			self.render("static/msg.html",msg="查询结果为空")
		
		
	elif frpr=="pr":
		rdpr=databasetabletodataframe2("insurancepremiumtable",rcuna)
		if len(rdpr)>0:
			rdpr=rdpr[rdpr.date>startd]
			rdpr=rdpr[rdpr.date<stopd]
	
			df=rdpr.reset_index(drop=True)
		
			l=len(df)

			if l>0:
				amount=round(rdpr.suas.sum(),2)
				ldf=df.index.tolist()
				self.render("static/rqprresult.html",df=df,amount=amount,ldf=ldf)
			else:
				self.render("static/msg.html",msg="查询结果为空")
		else:
			self.render("static/msg.html",msg="查询结果为空")
	else:
		self.render("static/msg.html",msg="数据库连结错误。")

	

#########################################################################################
#出货模块：1.表单数据获取：出货文件号，出货材料名称，包装规格，数量，出货仓库，业务员，助理名称。
#2.根据材料信息搜索库存并按有效期近到远排列，从上到下依次满足数量要求。如数量不足则弹出提示。
#3.如果批号数量正常则生成有效记录并存入库存表，并在末尾标识出货文件号以供储运更新库存。
#4.业务部根据库存表出货文件号生成出货单并加上客户信息，运费，保费，是否增加托盘记录。运保费预录入。等储运部确认。
#5.生成WORD文档供业务部下载打印及存档。

#根据材料，仓库，包装规格，数量出符合条件的库存汇总:DataFrame
###############################################################################################################################################
##选择拟出货的客户名称
class Selectcustomer(tornado.web.RequestHandler):
    @checkdelievery
    def get(self):	
	saonl=updateda(da)
	self.render("static/selectcustomer.html",cuna="请输入客户名称关键字",saonl=saonl)
    @checkdelievery
    def post(self):
        saon=self.get_argument('saon')
	cuna=self.get_argument('cuna')
	cunaq="%"+cuna+"%"
	sql="select cuna from customerinfotable where cuna like %s "
	sl=getsaleslist()
	al=getasisterlist()	
	try :
		rec=database.query(sql,cunaq)
		if len(rec)==1:
			rcuna=rec[0]["cuna"]

			self.render("static/selectcustomerconfirm.html",saon=saon,cuna=rcuna,sl=sl,al=al)
		elif len(rec)>1:
			l=len(rec)
			i=0
			lc=[]
			while i<l:
				lc.append(rec[i]["cuna"])
				i=i+1
			self.render("static/selectcustomerm.html",saon=saon,cuna=lc)
		else:
			self.render("static/selectcustomer.html",cuna="没有这种客户，请重新输入。")
	except:
		self.render("static/msg.html",msg="数据库连结失败。请联系管理员。")	

###############################################################################################################################################
#######记录出货文件号与相应客户名称及日期
class Saoncustomer(tornado.web.RequestHandler):
	
	def post(self):
        	saon=self.get_argument('saon')
		cuna=self.get_argument('cuna')
		wana=self.get_argument('wana')
		sape=self.get_argument("sape")
		mape=self.get_argument("mape")
		sql="insert into saoncustomertable (saon,cuna,date) values (%s,%s,%s)"
		sqlc="select saon from saoncustomertable where saon=%s"
		sqlcs="select * from customerinfotable where cuna=%s"
		recs=database.query(sqlcs,cuna)
		if recs:
		
			if not len(saon)==0:
				try:
					rec=database.query(sqlc,saon)
					if len(rec)>0:
						self.write("<center><h1>出货文件号重复。</h1></center>")
					elif len(rec)==0:
						database.execute(sql,saon,cuna,strtoday())
						self.render("static/saoncustomerc.html",saon=saon,wana=wana,sape=sape,mape=mape)
				except:
					self.render("static/msg.html",msg="数据库连结失败。请联系管理员。")
			else:
				self.render("static/msg.html",msg="请输入正确的出货文件号")
		else:self.render("static/msg.html",msg="没有这个客户")
#########################################################################################
#######出货接口

class Salesd(tornado.web.RequestHandler):
	@checkdelievery	
	def get(self):
		try:
			saon=self.get_argument("saon")
			wana=self.get_argument("wana")
			sape=self.get_argument("sape")
			mape=self.get_argument("mape")
			self.render("static/delieverym.html",saon=saon,wana=wana,sape=sape,mape=mape)
		except:
			
			self.render("static/selectcustomer.html")
	@checkdelievery	
	def post(self):
		saon=self.get_argument("saon")
		wana=self.get_argument("wana")
		sape=self.get_argument("sape")
		mape=self.get_argument("mape")
		kmana=self.get_argument("mana")
		mana="%"+kmana+"%"
		cuna=getcunafromsaon(saon)
		cmpi=getcustomerallprice(cuna)
		rec=database.query("select * from stocktable where wana=%s and mana like %s   order by mana desc,past asc, right(exda,4) asc,left(exda,2) asc ",wana,mana)
		if len(rec)>0:
			df=pd.DataFrame(rec)
			#df=df.sort_values(by='exda'.decode('utf-8'),ascending=True)
			ldf=df.index.tolist()
			self.render("static/delieverymb.html",saon=saon,wana=wana,sape=sape,mape=mape,ldf=ldf,df=df,cmpi=cmpi)
		else:
			self.render("static/msg.html",msg="无可用批号")
		
		

##################################################################################################################################################
class Delievery(tornado.web.RequestHandler):
	
     def post(self):
         banupast=self.get_argument('banu')
         banu=banupast.split(",")[0]
         past=float(banupast.split(",")[1])
         mana=banupast.split(",")[2]
         saon=self.get_argument('saon')
         quan=float(self.get_argument('quan'))
         wana=self.get_argument('wana')
         sape=self.get_argument('sape')
         mape=self.get_argument('mape')
         pric=self.get_argument('pric')
         insf=self.get_argument('insf')
         cuna=getcunafromsaon(saon)
         if pric=="":
             pirl=getrecentprice(cuna,mana)
             pric=pirl[1]
             insf=pirl[0]
         else:
             pass
         rec=recstockb(banu,past,quan,wana)
         if len(rec)>0:
             #self.write(str(rec.columns))
             writeintobedealwithtable(saon,sape,mape,rec,pric,insf)
             if prereducestock(rec):
                 df=pd.DataFrame(rec)
                 ldf=df.index.tolist()
                 pdddf=partdelieverydetail(saon)
                 lpdddf=pdddf.index.tolist()
                 self.render("static/saoncustomerct.html",saon=saon,wana=wana,sape=sape,mape=mape,df=df,ldf=ldf,pdddf=pdddf,lpdddf=lpdddf)
             else:
                self.render("static/msg.html",msg="库存预扣失败")
         else:
            self.render("static/msg.html",msg="批号错误.")


###############################################################################################################
####根据出货文件号生成出货单。
class Showdelieverydocument(tornado.web.RequestHandler):
    @checkdelievery
    def get(self):	
	self.render("static/toshowdelievery.html")
    @checkdelievery
    def post(self):
        saon=self.get_argument('saon')
	dname=self.get_secure_cookie('speedfairuser')
	cdd=createdelieverydocument(saon,dname)
	self.render("static/downloaddd.html")
###############################################################################################################################################


################################################################################################################################################
####出货单文件下载
class Downloadd(tornado.web.RequestHandler):
    #@checkdelievery
    def post(self):	

	docc=self.get_argument("docc")
	dusername=self.get_secure_cookie('speedfairuser')
	docname=dusername+docc
    	self.set_header ('Content-Type', 'application/octet-stream')
	hc='attachment; filename='+docname
    	self.set_header ('Content-Disposition', hc)

	dp="static/"+docname
    	with open(dp) as f:
        	while True:
            		data = f.read()
            		if not data:
               			 break
            		self.write(data)

    	self.finish()


#################################################################################################################################################
####计算运费的公开网页接口。支持匿名登陆使用
class Freightcalc(tornado.web.RequestHandler):
    def get(self):
	#self.write(str(self.request.headers))	
	self.render("static/freightcalc.html")
    def post(self):
        frar=self.get_argument('frar')
	frwe=self.get_argument('frwe')
	sfdl=self.get_argument("sfdl")
	sfcu=self.get_argument("sfcu")
	
	reslt=calcfreight(frar,frwe,sfdl,sfcu)
	if reslt==0:
		self.write("<center><h1>没有这个价档 </h1></center>")
	elif reslt==1:
		self.write("got gnormal")
	else:
		self.write("<script src='static/top.js'></script>  ")
		self.write("<p><center><h1>运费： ")
		self.write(str(reslt["frei"]))
		self.write(" 元。</h1></center></p>")
		self.write("<p><center><h1>客户地区： ")
		self.write(reslt["frar"])
		self.write("。</h1></center></p>")
		self.write("<p><center><h1>计费重量： ")
		self.write(str(reslt["frwe"]))
		self.write(" 吨。</h1></center></p>")
		self.write("<p><center><h1>运费单价： ")
		self.write(str(reslt["frpr"]))
		self.write(" 元/吨。</h1></center></p>")
		self.write("<p><center><h1> ")
		self.write(reslt["sfdl"])
		self.write("</h1></center></p>")
		self.write("<p><center><h1> ")
		self.write(reslt["sfcu"])
		self.write("</h1></center></p></html>")
###############################################################################################################################################
class Changefp(tornado.web.RequestHandler):
	@checkdelievery
	def get(self):
		self.render("static/changefreig.html")
	
	def post(self):
		switch=self.get_argument('frpr')
		saon=self.get_argument('saon')
		fpee=self.get_argument('fpee')
		if switch=="fr":
			sql="update freighttable set frei=%s where saon=%s"
			try:
				database.execute(sql,float(fpee),saon)
				self.render("static/msg.html",msg="运费已更新。")
			except:
				self.render("static/msg.html",msg="运费更新失败。可能是出货文件号错误。")
		elif switch=="pr":
			sql="update insurancepremiumtable set suas=%s where saon=%s"
			try:
				database.execute(sql,float(fpee),saon)
				self.render("static/msg.html",msg="保费已更新。")
			except:
				self.render("static/msg.html",msg="保费更新失败。可能是出货文件号错误。")
		else:
			self.render("static/msg.html",msg="未知错误，请联系管理员。")

############################################################################################################################################
class Calculatefreight(tornado.web.RequestHandler):
	@checkdelievery
	def get(self):
		self.render("static/calcfreight.html")

	
	def post(self):
		saon=self.get_argument('saon')
		sqlc="select * from saoncustomertable where saon=%s"
		sqlcs="select * from customerinfotable where cuna=%s"
		sqlm="select * from tobedealwithtable where saon=%s"
		dname=self.get_secure_cookie('speedfairuser')
		try:
			rec=database.query(sqlc,saon)
			
			if len(rec):
				
				cuna=rec[0]["cuna"]

				recs=database.query(sqlcs,cuna)
				#self.write(str(recs))
				frar=recs[0]["frar"]

				sfdl=recs[0]["sfdl"]

				sfcu=recs[0]["frsc"]

				tbcdd=tobeconfirmdelieverydataframe(saon)
				
				gwr=str(round(float((tbcdd["quan"]*(tbcdd["pagw"]+tbcdd["past"])).sum())/1000,2))
				
				if sfdl=="gdanger":
					gwr=str(getdfreightweight(float(gwr)))
				else:
					gwr=str(getnfreightweight(float(gwr)))


				#self.write(gwr)

				frei=calcfreight(frar,gwr,sfdl,sfcu)
				

				prsc=recs[0]["prsc"]
				rate=0.00045
				
				insf=tbcdd["insf"][0]
				if insf=="n":
					saam=getsaamn(tbcdd)
				else:
					saam=getsaam(tbcdd)
				
				suas=round(saam*rate,1)
				

				###write to freighttable and insurancepremiumtable
				
				#写入运费表
				try:
					sqls="select saon from freighttable where saon=%s"
					sql="insert into freighttable (saon,cuna,frwe,pric,sfcu,frei,lona,date) values (%s,%s,%s,%s,%s,%s,%s,%s)"
					rec=database.query(sqls,saon)
					
					if len(rec)==0:

						database.execute(sql,saon,cuna,gwr,frei["frpr"],sfcu,frei["frei"],"鹏远",strtoday())

						#self.render("static/msg.html",msg="运费计算完毕。")
					else:
						database.execute("delete from freighttable where saon=%s ",saon)
						database.execute(sql,saon,cuna,gwr,frei["frpr"],sfcu,frei["frei"],"鹏远",strtoday())
						#self.render("static/msg.html",msg="运费已重新计算")

				except:
					self.render("static/msg.html",msg="写入运费数据库失败。请联系管理员。")
				#写入保费表
				try:
					sqls="select saon from insurancepremiumtable where saon=%s"
					sql="insert into insurancepremiumtable (saon,cuna,suas,rate,prsc,date) values (%s,%s,%s,%s,%s,%s)"
					rec=database.query(sqls,saon)
					
					if len(rec)==0:

						database.execute(sql,saon,cuna,suas,rate,prsc,strtoday())
						#self.render("static/msg.html",msg="保费计算完毕。")
					else:
						database.execute("delete from insurancepremiumtable where saon=%s ",saon)
						
						database.execute(sql,saon,cuna,suas,rate,prsc,strtoday())
						
						#self.render("static/msg.html",msg="运保费已重新计算。")
						cdd=createdelieverydocument(saon,dname)
						
						self.render("static/downloaddd.html")
						
						
				except:
					self.render("static/msg.html",msg="写入保费数据库失败。请联系管理员.")
				cdd=createdelieverydocument(saon,dname)
				self.render("static/downloaddd.html")
				


				
			elif len(rec)==0:
				self.render("static/msg.html",msg="没有这个出货文件号.")
		except:
			self.render("static/msg.html",msg="数据库连结失败或者出货文件号错误。请联系管理员。")
###############################################################################################################################################
class Rqdg(tornado.web.RequestHandler):
    	@authenticated
	def get(self):
		self.render("static/rqdg.html")
    	@authenticated
	def post(self):

		#dstart=self.get_argument("dstart")
        	startd=self.get_argument('startd')
		stopd=self.get_argument('stopd')
		rqkw=self.get_argument("rqkw")
		rqop=self.get_argument("rqop")
		#dstop=self.get_argument("dstop")
		#self.write(startd)
		#self.write(stopd)
		#self.write(rqkw)
		#self.write(rqop)
		try:
			startd=startd.split("-")[0]+startd.split("-")[1]+startd.split("-")[2]
			stopd=stopd.split("-")[0]+stopd.split("-")[1]+stopd.split("-")[2]
		except:
			startd=strtoday()[0:5]+"0101"
			stopd=str(int(strtoday()[0:5])+1)+"0101"
		if rqop=="cuna":
			kcuna="%"+rqkw+"%"
			rec=database.query("select * from delieverytable where cuna like %s and owda > %s and owda < %s",kcuna,startd,stopd)
			if len(rec)>0:
				df=pd.DataFrame(rec)
				amount=(df["past"]*df["quan"]).sum()
				df=df.set_index("id")
				df=df.reset_index(drop=True)
				ldf=df.index.tolist()
				self.render("static/rqdgresult.html",amount=amount,df=df,ldf=ldf)
			else:
				self.render("static/msg.html",msg="查询结果为空")
		elif rqop=="banu":
			kbanu="%"+rqkw+"%"
			rec=database.query("select * from delieverytable where banu like %s and owda > %s and owda < %s",kbanu,startd,stopd)
			if len(rec)>0:
				df=pd.DataFrame(rec)
				amount=(df["past"]*df["quan"]).sum()

				df=df.set_index("id")
				df=df.reset_index(drop=True)
				ldf=df.index.tolist()
				
				self.render("static/rqdgresult.html",amount=amount,df=df,ldf=ldf)
			else:
				self.render("static/msg.html",msg="查询结果为空")
		elif rqop=="mana":
			kmana="%"+rqkw+"%"
			rec=database.query("select * from delieverytable where mana like %s and owda > %s and owda < %s",kmana,startd,stopd)
			if len(rec)>0:
				df=pd.DataFrame(rec)
				amount=(df["past"]*df["quan"]).sum()
				df=df.set_index("id")
				df=df.reset_index(drop=True)
				ldf=df.index.tolist()
				self.render("static/rqdgresult.html",amount=amount,df=df,ldf=ldf)
			else:
				self.render("static/msg.html",msg="查询结果为空")
		elif rqop=="sape":
			ksape="%"+rqkw+"%"
			rec=database.query("select * from delieverytable where sape like %s and owda > %s and owda < %s",ksape,startd,stopd)
			if len(rec)>0:
				df=pd.DataFrame(rec)
				amount=(df["past"]*df["quan"]).sum()
				df=df.set_index("id")
				df=df.reset_index(drop=True)
				ldf=df.index.tolist()
				self.render("static/rqdgresult.html",amount=amount,df=df,ldf=ldf)
			else:
				self.render("static/msg.html",msg="查询结果为空")
		elif rqop=="incl":
			kincl="%"+rqkw+"%"
			rec=database.query("select * from delieverytable where sape like %s and owda > %s and owda < %s",kincl,startd,stopd)
			if len(rec)>0:
				df=pd.DataFrame(rec)
				amount=(df["past"]*df["quan"]).sum()
				df=df.set_index("id")
				df=df.reset_index(drop=True)
				ldf=df.index.tolist()
				self.render("static/rqdgresult.html",amount=amount,df=df,ldf=ldf)
			else:
				self.render("static/msg.html",msg="查询结果为空")
		else:
			pass

###############################################################################################################################################
class Rqpi(tornado.web.RequestHandler):
    	@authenticated
	def get(self):
		self.render("static/rqpi.html")
    	@authenticated
	def post(self):
        	startd=self.get_argument('startd')
		stopd=self.get_argument('stopd')
		rqkw=self.get_argument("rqkw")
		rqop=self.get_argument("rqop")
		try:
			startd=startd.split("-")[0]+startd.split("-")[1]+startd.split("-")[2]
			stopd=stopd.split("-")[0]+stopd.split("-")[1]+stopd.split("-")[2]
		except:
			startd=strtoday()[0:5]+"0101"
			stopd=str(int(strtoday()[0:5])+1)+"0101"
		if rqop=="puon":
			kpinu="%"+rqkw+"%"
			rec=database.query("select * from putinwhtable where puon like %s and ewda > %s and ewda < %s",kpinu,startd,stopd)
			if len(rec)>0:
				df=pd.DataFrame(rec)
				amount=(df["past"]*df["quan"]).sum()
				df=df.set_index("id")
				df=df.reset_index(drop=True)
				ldf=df.index.tolist()
				self.render("static/rqpiresult.html",df=df,ldf=ldf)
			else:
				self.render("static/msg.html",msg="查询结果为空")
		elif rqop=="banu":
			kbanu="%"+rqkw+"%"
			rec=database.query("select * from putinwhtable where banu like %s and ewda > %s and ewda < %s",kbanu,startd,stopd)
			if len(rec)>0:
				df=pd.DataFrame(rec)
				amount=(df["past"]*df["quan"]).sum()

				df=df.set_index("id")
				df=df.reset_index(drop=True)
				ldf=df.index.tolist()
				
				self.render("static/rqpiresult.html",df=df,ldf=ldf)
			else:
				self.render("static/msg.html",msg="查询结果为空")
		elif rqop=="mana":
			kmana="%"+rqkw+"%"
			rec=database.query("select * from putinwhtable where mana like %s and ewda > %s and ewda < %s",kmana,startd,stopd)
			if len(rec)>0:
				df=pd.DataFrame(rec)
				amount=(df["past"]*df["quan"]).sum()
				df=df.set_index("id")
				df=df.reset_index(drop=True)
				ldf=df.index.tolist()
				self.render("static/rqpiresult.html",df=df,ldf=ldf)
			else:
				self.render("static/msg.html",msg="查询结果为空")
		else:
			pass


###############################################################################################################################################
class Confirmdelievery(tornado.web.RequestHandler):
	@checkputinwh
	def get(self):
		try:
			tl=getsaonlist(getinfotobedealwithtable())
			self.render('static/tbdw.html',tl=tl)
		except:
			self.render("static/msg.html",msg="没有出货单")
	
	def post(self):
		saon=self.get_argument("saon")
		acco=self.get_argument("acco")
		if acco=="conf":
			rst=confirmdelievery(saon)
			if rst:
				database.execute("delete from tobedealwithtable where saon=%s",saon)
				self.render("static/msg.html",msg="出货完毕")
			else:
				self.render("static/msg.html",msg="出货失败，请联系管理员")


		elif acco=="undo":
			rst=undodelievery(saon)
			if rst:
				database.execute("delete from tobedealwithtable where saon=%s",saon)				
				self.render("static/msg.html",msg="撤消出货完毕")
			else:
				self.render("static/msg.html",msg="撤消出货失败，请联系管理员")
		else:
			pass

#########################################################################################
#######进仓接口

class Putinwh(tornado.web.RequestHandler):
	@checkputinwh
	def get(self):
		try:
			spinu=self.get_argument("pinu")
			puon=self.get_argument("ponu")
			wana=self.get_argument("wana")
			l=[]
			l.append(spinu)
			lw=[]
			lw.append(wana)
			self.render("static/putinwh.html",spinu=l,puon=puon,swana=lw)
		except:
			
			spinu=updateia(ia)
			lw=["厦门","上海","北京","广州"]
			self.render("static/putinwh.html",spinu=spinu,puon="",swana=lw)
	@checkputinwh
	def post(self):
		
		mana=self.get_argument("mana")
		banu=self.get_argument("banu")
		exda=self.get_argument("exda")
		past=self.get_argument("past")
		pagw=self.get_argument("pagw")
		quan=self.get_argument("quan")
		pinu=self.get_argument("pinu")
		wana=self.get_argument("wana")
		puon=self.get_argument("ponu")
		#ewda=strtoday()
		pipe=self.get_secure_cookie('speedfairuser')
		putinwh(mana,banu,exda,past,pagw,quan,pinu,wana,puon,pipe)
		self.render("static/msgcps.html",pinu=pinu,puon=puon,wana=wana)
#########################################################################################

			
###############################################################################################################
####根据进货文件号生成进货单。
class Showinputdoc(tornado.web.RequestHandler):
    #@checkdelievery
    def get(self):	
	self.render("static/toshowinputdoc.html")
    #@checkdelievery
    def post(self):
        pinu=self.get_argument('pinu')
	rec=database.query("select * from putinwhtable where pinu=%s",pinu)
	if len(rec)>0:
		df=pd.DataFrame(rec)
		df=df.set_index("id")
		df=df.reset_index(drop=True)
		pinu=df.ix[0]["pinu"]
		puon=df.ix[0]["puon"]
		ewda=df.ix[0]["ewda"]
		dictinfo={"pinu":pinu,"puon":puon,"ewda":ewda}
		document=docx.Document("putindocument.docx")
		#document.tables[0].style='Table Grid'
		loopreplaceindocx(document,dictinfo)
		i=0
		l=len(df)
		while i<l:

			document.tables[0].cell(i+1,0).text=df.ix[i]["mana"]
			document.tables[0].cell(i+1,1).text=str(df.ix[i]["past"])+"公斤/桶".decode('utf8')
			document.tables[0].cell(i+1,2).text=str(df.ix[i]["quan"])+"桶".decode('utf8')
			document.tables[0].cell(i+1,3).text=str(df.ix[i]["past"]*df.ix[i]["quan"])+"公斤".decode('utf8')
			document.tables[0].cell(i+1,4).text=df.ix[i]["banu"]
			document.tables[0].cell(i+1,5).text=df.ix[i]["exda"]
			
			i=i+1
			document.tables[0].add_row()
		document.tables[0].cell(l+1,0).text="总计：".decode('utf8')
		document.tables[0].cell(l+1,1).text=""
		document.tables[0].cell(l+1,2).text=str(df["quan"].sum())+"桶".decode('utf8')
		document.tables[0].cell(l+1,3).text=str((df["past"]*df["quan"]).sum())+"公斤".decode('utf8')
		document.tables[0].cell(l+1,4).text=""

		pname=self.get_secure_cookie('speedfairuser')
		pidcdir="static/"+pname+"pidc.docx"
		#document.tables[0]._tblPr.set='tblBorders'
		document.save(pidcdir)
		self.render("static/putinwhdoc.html")
	else:self.render("static/msg.html",msg="没有这个进货号")





	
		
#########################################################################################

class Currentstockdetail(tornado.web.RequestHandler):
	def get(self):	
		csd=detailstock()
		self.render("static/detailstock.html",csd=csd)	
			
		
		



#########################################################################################
###有效期4个月以内的库存查询
class Querystockexda(tornado.web.RequestHandler):
    @authenticated
    def get(self):	
	rec=exdarq(exdashow(4))
	if len(rec)>0:
		dfw=pd.DataFrame(rec)
		df=dfw.reset_index(drop=True)
		quantity=(df.quan*df.past).sum()
		amount=df.quan.sum()
		ldf=df.index.tolist()
		self.render("static/rqstockresultexda.html",df=df,quantity=quantity,amount=amount,ldf=ldf)
	else:
		self.render("static/msg.html",msg="无"+exdashow(7)+str(len(rec)))
#################################################################################################################
###客户信息录入,如果发现客户名称重复的话,直接更新.
class Customerinfoinput(tornado.web.RequestHandler):
	@checkdelievery
	def get(self):
		areas=getfrarall()
		self.render("static/customerinfoinput.html",areas=areas)

	def post(self):
		cuna=self.get_argument("cuna")

		repl=self.get_argument("repl")

		inad=self.get_argument("inad")

		cosi=self.get_argument("cosi")

		ctel=self.get_argument("ctel")

		cmob=self.get_argument("cmob")

		cfax=self.get_argument("cfax")

		cpos=self.get_argument("cpos")

		frsc=self.get_argument("frsc")
		prsc=self.get_argument("prsc")
		tanu=self.get_argument("tanu")

		rema=self.get_argument("rema")
		if rema=="":
			rema="空"
		else:
			pass
		pall=self.get_argument("pall")
		incl=self.get_argument("incl")
		sfdl=self.get_argument("sfdl")
		frar=self.get_argument("frar")
		
		sqlci="insert into customerinfotable ( cuna,repl,inad,cosi,ctel,cmob,cfax,cpos,frsc,prsc,tanu,rema,pall,incl,sfdl,frar) values (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
		sqldci="delete from customerinfotable where cuna=%s"
		rec=database.query("select * from customerinfotable where cuna=%s ",cuna)
		if len(rec)>0:
			database.execute(sqldci,cuna)
			database.execute(sqlci,cuna,repl,inad,cosi,ctel,cmob,cfax,cpos,frsc,prsc,tanu,rema,pall,incl,sfdl,frar)
			#self.render("static/msg.html",msg="客户信息更新完毕")

		else:
			database.execute(sqlci,cuna,repl,inad,cosi,ctel,cmob,cfax,cpos,frsc,prsc,tanu,rema,pall,incl,sfdl,frar)
			self.render("static/msg.html",msg="客户信息输入完毕")

#################################################################################################################
###modify tobedealwithtable infomation
class Modifytobedealwith(tornado.web.RequestHandler):
    @checkdelievery
    def get(self):
        try:
            tl=getsaonlist(getinfotobedealwithtable())
            self.render('static/modifytbdw.html',tl=tl)
        except:
            self.render("static/msg.html",msg="没有出货单")

    def post(self):	
        saon=self.get_argument("saon")
        rs=getonetobedealwithtable(saon)
        ldf=rs.index.tolist()
        self.render('static/deloneintobedealwith.html',choice=rs,saon=saon,ldf=ldf)
        
         
#################################################################################################################
###modify tobedealwithtable infomation
class Modifytobedealwith2(tornado.web.RequestHandler):
    @checkdelievery    

    def post(self):	
        saon=self.get_argument("saon")
        bqsw=self.get_argument("choice")
        bqswl=bqsw.split(",")

       
        

        if undoonedelievery(bqswl,saon):
            try:
                tl=getsaonlist(getinfotobedealwithtable())
                self.render('static/modifytbdw.html',tl=tl)
            except:
                self.render("static/msg.html",msg="没有出货单")
        else:
            self.render("static/msg.html",msg="some wrong happen,please contact with administrator")
#####生成出货文件号生成结果rl[0]是可能的发货号列表，rl[1]是新的a保持a是最后一个号特征
def updateda(a):
	global da
	if a[0:6]==strtoday()[0:6]:
		ac=str(int(a[6:9])+1)
		l=len(ac)
		if l==1:
			d="00"+ac
		elif l==2:
			d="0"+ac
		elif l==3:
			d=ac
		else:
			pass
	else:
		d="001"
	na=strtoday()[0:6]+d
	nas=na[2:9]
	nh=nas[0:2]
	nt=nas[2:7]
	l=[]
	l.append("G"+nh+"B"+nt)
	l.append("S"+nh+"B"+nt)
	l.append("B"+nh+"B"+nt)
	l.append("X"+nh+"B"+nt)
	l.append("D"+nh+"B"+nt)
	da=na
	return l
#########################################################################################
#####生成进货文件号生成结果rl[0]是可能的发货号列表，rl[1]是新的a保持a是最后一个号特征
def updateia(a):
	global ia
	if a[0:6]==strtoday()[0:6]:
		ac=str(int(a[6:9])+1)
		l=len(ac)
		if l==1:
			d="00"+ac
		elif l==2:
			d="0"+ac
		elif l==3:
			d=ac
		else:
			pass
	else:
		d="001"
	na=strtoday()[0:6]+d
	nas=na[2:9]
	nh=nas[0:2]
	nt=nas[2:7]
	l=[]
	l.append("X"+nh+"B"+nt)
	l.append("S"+nh+"B"+nt)
	l.append("B"+nh+"B"+nt)
	l.append("G"+nh+"B"+nt)
	ia=na
	return l
###############根据仓库名称返回进货文件号并更新全局进货文件号变量
def getnewiaandupdateia(wana):
	global ia
	lw=['厦门','上海','北京','广州']
	li=updateia(ia)
	return li[lw.index(wana)]


###############
###############批量导入进货单
def dropduplicates(d,a):
	result=d[a].drop_duplicates().tolist()
	return result
def getpacklistandputinwh(pipe):
	df=pd.read_excel("packlist/packlist.xlsx")
	dwl=dropduplicates(df,'wana')
	for dwle in dwl:
		dfwa=df[df['wana']==dwle]
		dfwal=dropduplicates(dfwa,'puno')
		for dfwap in dfwal:
			onepunopl= dfwa[dfwa['puno']==dfwap].reset_index(drop=True)[['puno','mana','banu','exda','past','pagw','quan','wana']]
			lp=0
			lia=getnewiaandupdateia(dwle.encode('utf-8'     )) 
			while lp<len(onepunopl):
				oppl=onepunopl.ix[lp]
				#print lia,oppl['puno'],oppl['mana'],oppl['banu'],oppl['past'],oppl['pagw'],oppl['drums'],oppl['wana']
				putinwh(oppl['mana'],oppl['banu'],oppl['exda'],oppl['past'],oppl['pagw'],oppl['quan'],lia,oppl['wana'],oppl['puno'],pipe)
				lp=lp+1






#############################
###撤单，撤去已发货记录，还回库存
class Undo(tornado.web.RequestHandler):
    @checkadmin
    def get(self):
        self.render("static/undo.html")
    @checkadmin
    def post(self):
        saon=self.get_argument("saon")
        undoandtrans.undodelievery(saon)
        self.render("static/msg.html",msg='已撤')
###撤单，撤去已发货记录，还回库存
class Trans(tornado.web.RequestHandler):
    @checkadmin
    def get(self):
        self.render("static/trans.html")
    @checkadmin
    def post(self):
        banu=self.get_argument("banu")
        quan=int(self.get_argument("quan"))
        oldwana=self.get_argument("oldwana")
        newwana=self.get_argument("newwana")
        undoandtrans.transbanu(banu,quan,oldwana,newwana)
        self.render("static/msg.html",msg=('转库完毕') )  
class UploadFileHandler(tornado.web.RequestHandler):
    
    @checkadmin
    def get(self):
        self.render("static/uploadfile.html")

    def post(self):
        #文件的暂存路径
        upload_path=os.path.join(os.path.dirname(__file__),'packlist')  
        self.write("今天进库清单")
        #提取表单中‘name’为‘file’的文件元数据
        file_metas=self.request.files['file']
        for meta in file_metas:
            filename=meta['filename']
            filepath=os.path.join(upload_path,filename)
            #有些文件需要已二进制的形式存储，实际中可以更改
            with open(filepath,'wb') as up:      
                up.write(meta['body'])
            self.write('上传完毕!')
	pipe=self.get_secure_cookie('speedfairuser')
	getpacklistandputinwh(pipe)
	self.write('库存批量导入完毕')

#########################################################################################

application = tornado.web.Application([
    (r"/", MainHandler),
    (r"/signin", SigninHandler),
    (r"/addnewuser",Addnewuser),
    (r"/addnewuserp",Addnewuserp),
    (r"/viewuser",Viewuser),
    (r"/changepwd",Changepwd),
    (r"/deleteuser",Deleteuser),
    (r"/toexcel",Toexcel),
    (r"/querystock",Querystock),
    (r"/querystockexda",Querystockexda),
    (r"/rqdg",Rqdg),
    (r"/rqpi",Rqpi),
    (r"/salesd",Salesd),
    (r"/delievery",Delievery),
    (r"/freightcalc",Freightcalc),
    (r"/selectcustomer",Selectcustomer),
    (r"/saoncustomer",Saoncustomer),
    (r"/calculatefreight",Calculatefreight),
    (r"/rqfrpr",Rqfrpr),
    (r"/showdelieverydocument",Showdelieverydocument),
    (r"/downloadd",Downloadd),
    (r"/changefp",Changefp),
    (r"/confirmdelievery",Confirmdelievery),
    (r"/putinwh",Putinwh),
    (r"/showinputdoc",Showinputdoc),
    (r"/customerinfoinput",Customerinfoinput),
    (r"/currentstockdetail",Currentstockdetail),
    (r"/setiada",Setiada),
    (r"/modifytobedealwith",Modifytobedealwith),
    (r"/modifytobedealwith2",Modifytobedealwith2),
    (r"/undo",Undo),
    (r"/trans",Trans),
    (r"/uploadfile",UploadFileHandler)
],cookie_secret="61oETzKXQAGaYdkL5gEmGeJJFuYh7EQnp2XdTP1o/Vo=jjkkd",static_path=os.path.join(os.path.dirname(__file__), "static"))

if __name__ == "__main__":
    application.listen(8080)
    tornado.ioloop.IOLoop.instance().start()
