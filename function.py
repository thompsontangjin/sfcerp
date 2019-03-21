# -*- coding: utf-8 -*-
import torndb
import pandas as pd
import datetime
import docx
import os
from database import database
from authenticate import *



def createdelieverydocument(saon,dname):


	##计算托盘数量


	if len(saon)>0:
		###got saon
		sqlt="select * from tobedealwithtable where saon=%s"
		rect=database.query(sqlt,saon)
		sqld="select * from delieverytable where saon=%s"
		recd=database.query(sqld,saon)
		if len(rect)>0:
			rec=rect
		elif len(recd)>0:
			rec=recd
		else:
			rec=''

		if len(rec)>0:
			df=pd.DataFrame(rec)
			df=df.set_index("id")
			df=df.reset_index(drop=True)
			######要求一行一行相乘取小数点后两位后相加
			#saam=(df["past"]*df["quan"]*df["sapr"]).sum()
			#saam=getsaam(df)
			sumquan=str(df["quan"].sum())
			sumqulity=str((df["quan"]*df["past"]).sum())
			i=0

			sape=df.ix[0]["sape"]
			insf=df.ix[0]["insf"]
			if insf=="n":
				saam=getsaamn(df)
			else:
				saam=getsaam(df)
			###got cuna  date
			sqlc="select * from saoncustomertable where saon=%s"
			try:
				recc=database.query(sqlc,saon)
				cuna=recc[0]["cuna"]
				date=recc[0]["date"]
				
			except:
				l=[]
				l.append("数据库无法连结")
				return l
			sqlci="select * from customerinfotable where cuna=%s"
			try:
				recci=database.query(sqlci,cuna)
				repl=recci[0]["repl"]
				cosi=recci[0]["cosi"]
				cpos=recci[0]["cpos"]
				try:
					if len(cpos)>0:
						pass
					else:
						pass
				except:
					cpos="无".decode('utf8')
				cfax=recci[0]["cfax"]
				ctel=recci[0]["ctel"]
				if len(ctel)==0:
					ctel="无"
				else:
					pass
				cmob=recci[0]["cmob"]
				tanu=recci[0]["tanu"]
				remark=recci[0]["rema"]
				try:
					if len(remark)>0:
						pass
					else:
						pass
				except:
					remark="".decode('utf8')
				
			except:
				l=[]
				l.append("数据库无法连结")
				return l
			sqlf="select * from freighttable where saon=%s"
			try:
				recf=database.query(sqlf,saon)
				pric=recf[0]["pric"]
				frwe=recf[0]["frwe"]
				frei=recf[0]["frei"]
				sfcu=recf[0]["sfcu"]
				if sfcu=="Z":
					frcu="自付"
				else:
					frcu="代垫"
				
			except:
				l=[]
				l.append("数据库无法连结")
				return l
			sqli="select * from insurancepremiumtable where saon=%s"
			try:
				reci=database.query(sqli,saon)
				suas=reci[0]["suas"]
				prsc=reci[0]["prsc"]
				if prsc=="Z":
					prcu="自付"
				else:
					prcu="代垫"

				
			except:
				l=[]
				l.append("数据库无法连结")
				return l
				
			if prsc=="D":
				amount=saam+suas
			else:
				amount=saam
			if sfcu=="D":
				amount=amount+frei
			else:
				amount=amount
			######pallet
			if len(cuna)>0:
				try:
					recpallet=database.query("select pall from customerinfotable where cuna=%s",cuna)
					palletif=recpallet[0]["pall"]
					if palletif=="y":
						
						palletsamount=str(pallets(saon))
						
						palletinfo="225公斤大桶请务必打托盘出货，托盘数量：".decode('utf8')+palletsamount
						
					else :
						palletinfo=""
				except:
					pass
			else:
				pass
			if len(palletinfo)>0:
				remark=remark+palletinfo
			else:
				pass


			dictinfo={"saon":saon,"sape":sape,"cuna":cuna,"date":date,"repl":repl,"cosi":cosi,"cpos":cpos,"ctel":ctel,"cfax":cfax,"cmob":cmob,"suas":str(suas),"prcu":prcu.decode('utf8'),"frwe":str(frwe),"pric":str(pric),"frei":str(frei),"frcu":frcu.decode('utf8'),"amount":str(amount),"tanu":tanu,"remark":remark,"saam":str(saam)}


			if ifcumana(cuna):
				document=docx.Document("delieverydocument.docx")
				loopreplaceindocx(document,dictinfo)
				writedatatodocxtable(document,df)
				dusername=dname
				docname="static/"+dusername+"dldc.docx"
				#document.tables[0]._tblPr.set='tblBorders'
				document.save(docname)
			else:
				dc=dictcumana(cuna)
				if len(dc)>0:
					document=docx.Document("delieverydocument2.docx")
					loopreplaceindocx(document,dictinfo)
					writedatatodocxtablecumana(document,df,dc)
					dusername=dname
					docname="static/"+dusername+"dldc.docx"
					#document.tables[0]._tblPr.set='tblBorders'
					document.save(docname)
				else:
					pass
			return dictinfo
			

		else:
			l=[]
			l.append("没有这个出货号")
			return l
	else:
		l=[]
		l.append("出货号不能为空")
		return l
################################################################################################################################################
#######计算含税单价货值,逐行计算后取小数点后两位有效数字后累加和.
def getsaam(df):
	l=len(df)
	i=0
	saam=0
	while i<l:
		saamcol=round(df.ix[i]["past"]*df.ix[i]["quan"]*float(df.ix[i]["sapr"]),2)
		saam=saam+saamcol
		i=i+1
	return saam
################################################################################################################################################
#######计算不含税单价货值,逐行计算后取小数点后两位有效数字后累加和.有待改进：同一种材料汇总数量算出不含税金额后取两位有效数字再算出含税金额再取两位有效数字后汇总。先对数据检查是否有同一个材料名称，如果有的话对同一材料重量汇总后放在第一行的数量更新(考虑生成一个新的数据列表或dict可能更好，然后对dict里面的材料名称和重量进行累计算和)，如果有同一材料不同包装也要汇总计算。
def getsaamn(df):
	df=washdata(df)
	listmana=dropduplicates(df,"mana")
	l=len(df)
	saam=0
	for lm in listmana:
		dff=df[df.mana==lm].reset_index(drop=True)
		saam=saam+round(round((dff.past*dff.quan).sum()*float(dff.sapr[0]),2)*1.16,2)
	return saam

####对数据中单价为0的对应数量全部清洗删除所在行并对数据索引重置,以方便对同一材料不同批号汇总>    重量进行不含税计算
def washdata(df):
	df=df[df.sapr!=u'0']
	return df
################################################################################################################################################
#####确认是否需要带物料号出货单
def ifcumana(cuna):
	sql="select * from cumanatable where cuna=%s"
	rec=database.query(sql,cuna)
	if len(rec)>0:
		return False
	else:
		return True
################################################################################################################################################
####取得物料号对应表
def dictcumana(cuna):
	sql="select * from cumanatable where cuna=%s"
	rec=database.query(sql,cuna)
	if len(rec)>0:
		df=pd.DataFrame(rec)
		df=df.reset_index(drop=True)
		l=len(rec)
		i=0
		dictcumana={}
		while i<l:
			dictcumana[df.ix[i]['mana']]=df.ix[i]['cumana']
			i=i+1
		return dictcumana
	else:
		return " "
#################################################################################################################
def getsaleslist():
	sqls="select username from namepwd where level=%s"
	rsl=[]
	recn=database.query(sqls,"sa")
	for recno in recn:
		rsl.append(recno)
	return rsl	
#################################################################################################################
def getasisterlist():
	sqls="select username from namepwd where level=%s"
	rsl=[]
	recn=database.query(sqls,"as")
	for recno in recn:
		rsl.append(recno)
	return rsl	







#################################################################################################################
def putinwh(mana,banu,exda,past,pagw,quan,pinu,wana,puon,pipe):

		ewda=strtoday()
		sqlst="insert into stocktable (puon,mana,banu,exda,past,pagw,quan,wana,ewda,pana) values (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
		sqlpi="insert into putinwhtable (puon,mana,banu,exda,past,pagw,quan,pinu,wana,ewda,pipe) values (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
		sqlu="update stocktable set quan=%s,ewda=%s where banu=%s and wana=%s"
		
		database.execute(sqlpi,puon,mana,banu,exda,past,pagw,quan,pinu,wana,ewda,pipe)
		rec=database.query("select * from stocktable where banu=%s and wana=%s   order by right(exda,4) asc ,left(exda,2) asc",banu,wana)
		if len(rec)>0:
			
			squan=float(quan)+rec[0]["quan"]
			database.execute(sqlu,squan,ewda,banu,wana)

		else:
			database.execute(sqlst,puon,mana,banu,exda,past,pagw,quan,wana,ewda,"公斤/桶".decode('utf8'))

			
		
#####################################################################################################################################
##########
def writedatatodocxtable(document,df):
	i=0
	l=len(df)
	while i<l:

		document.tables[0].cell(i+1,0).text=df.ix[i]["mana"]
		document.tables[0].cell(i+1,1).text=str(df.ix[i]["past"])
		document.tables[0].cell(i+1,2).text="公斤/桶".decode('utf8')
		document.tables[0].cell(i+1,3).text=str(df.ix[i]["quan"])
		document.tables[0].cell(i+1,4).text="桶".decode('utf8')
		document.tables[0].cell(i+1,5).text=str(df.ix[i]["past"]*df.ix[i]["quan"])
		document.tables[0].cell(i+1,6).text="公斤".decode('utf8')
		document.tables[0].cell(i+1,7).text=df.ix[i]["banu"]
		document.tables[0].cell(i+1,8).text=df.ix[i]["exda"][0:10]
		i=i+1
		document.tables[0].add_row().cells
	document.tables[0].cell(l+1,0).text="总计：".decode('utf8')
	document.tables[0].cell(l+1,1).text=""
	document.tables[0].cell(l+1,2).text=""
	document.tables[0].cell(l+1,3).text=str(df["quan"].sum())
	document.tables[0].cell(l+1,4).text="桶".decode('utf8')
	document.tables[0].cell(l+1,5).text=str((df["past"]*df["quan"]).sum())
	document.tables[0].cell(l+1,6).text="公斤".decode('utf8')
	document.tables[0].cell(l+1,7).text=""
	document.tables[0].cell(l+1,8).text=""

#####################################################################################################################################
##########
def writedatatodocxtablecumana(document,df,dc):
	i=0
	l=len(df)

	while i<l:

		document.tables[0].cell(i+1,0).text=df.ix[i]["mana"]
		document.tables[0].cell(i+1,1).text=dc[df.ix[i]["mana"]]
		document.tables[0].cell(i+1,2).text=str(df.ix[i]["past"])
		document.tables[0].cell(i+1,3).text="公斤/桶".decode('utf8')
		document.tables[0].cell(i+1,4).text=str(df.ix[i]["quan"])
		document.tables[0].cell(i+1,5).text="桶".decode('utf8')
		document.tables[0].cell(i+1,6).text=str(df.ix[i]["past"]*df.ix[i]["quan"])
		document.tables[0].cell(i+1,7).text="公斤".decode('utf8')
		document.tables[0].cell(i+1,8).text=df.ix[i]["banu"]
		document.tables[0].cell(i+1,9).text=df.ix[i]["exda"][0:10]
		i=i+1
		document.tables[0].add_row().cells
	document.tables[0].cell(l+1,0).text="总计：".decode('utf8')
	document.tables[0].cell(l+1,1).text=""
	document.tables[0].cell(l+1,2).text=""
	document.tables[0].cell(l+1,3).text=""
	document.tables[0].cell(l+1,4).text=str(df["quan"].sum())
	document.tables[0].cell(l+1,5).text="桶".decode('utf8')
	document.tables[0].cell(l+1,6).text=str((df["past"]*df["quan"]).sum())
	document.tables[0].cell(l+1,7).text="公斤".decode('utf8')
	document.tables[0].cell(l+1,8).text=""
	document.tables[0].cell(l+1,9).text=""

	

		
			
###############################################################################################################################################
def confirmdelievery(saon):
	try:
		cuna=database.query("select cuna from saoncustomertable where saon=%s",saon)[0]['cuna']
		ds=getonetobedealwithtable(saon)
		l=len(ds)
		i=0
		while i<l:
			#ds.ix[i]['banu']   ds.ix[i]['quan']  ds.ix[i]['wana']
			sqld="delete from stocktable where banu=%s and wana=%s"
			sqlu="update stocktable set quan=%s where banu=%s and wana=%s"
			sqlq="select quan from stocktable where banu=%s and wana=%s"
			sqli="insert into delieverytable (saon,cuna,mana,banu,exda,past,quan,wana,sape,mape,sapr,owda,insf) values (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
			try:
				cq=database.query(sqlq,ds.ix[i]['banu'],ds.ix[i]['wana'])[0]['quan']
				if cq==0:
				    database.execute(sqld,ds.ix[i]['banu'],ds.ix[i]['wana'])
				else:
				    pass
			except:
				pass

			###write to delieverytable
			
			database.execute(sqli,saon,cuna,ds.ix[i]['mana'],ds.ix[i]['banu'],ds.ix[i]['exda'],ds.ix[i]['past'],ds.ix[i]['quan'],ds.ix[i]['wana'],ds.ix[i]['sape'],ds.ix[i]['mape'],ds.ix[i]['sapr'],strtoday(),ds.ix[i]['insf'])
			
			i=i+1
		
		return True
	except:
		return False




###############################################################################################################################################有待改进。当撤消的批号不存在库存记录时会出错，这种情况发生在同时出同一库的同一个批号且刚好先出的已确认会导致库存表没有这个批号记录。
def undodelievery(saon):
	ds=getonetobedealwithtable(saon)
	l=len(ds)
	i=0
	ewda=strtoday()
	try:
		while i<l:
			#ds.ix[i]['banu']   ds.ix[i]['quan']  ds.ix[i]['wana']
			sqlu="update stocktable set quan=%s where banu=%s and wana=%s"
			sqlq="select quan from stocktable where banu=%s and wana=%s"
			sqli="insert into stocktable (wana,mana,banu,exda,past,quan,pagw,ewda,pana) values (%s,%s,%s,%s,%s,%s,%s,%s,%s)"

			try:
				cq=database.query(sqlq,ds.ix[i]['banu'],ds.ix[i]['wana'])[0]['quan']
				if cq==0:
					database.execute(sqlu,ds.ix[i]['quan'],ds.ix[i]['banu'],ds.ix[i]['wana'])
				else:
					cq=cq+ds.ix[i]['quan']
					database.execute(sqlu,cq,ds.ix[i]['banu'],ds.ix[i]['wana'])
			except:
				database.execute(sqli,ds.ix[i]['wana'],ds.ix[i]['mana'],ds.ix[i]['banu'],ds.ix[i]['exda'],ds.ix[i]['past'],ds.ix[i]['quan'],ds.ix[i]['pagw'],ewda,'公斤/桶'.decode('utf8'))
			i=i+1
		database.execute("delete from tobedealwithtable where saon=%s",saon)
		database.execute("delete from insurancepremiumtable where saon=%s",saon)
		database.execute("delete from freighttable where saon=%s",saon)
		database.execute("delete from saoncustomertable where saon=%s",saon)
		return True
	except:
		return False
###############################################################################################################################################
#delete one recorder in tobedealwithtable ,bqsw is banu,quan,sapr,wana,in a list.
def undoonedelievery(bqswl,saon):
    try:
        sqlu="update stocktable set quan=%s where banu=%s and wana=%s"
        sqlq="select quan from stocktable where banu=%s and wana=%s"
        cq=database.query(sqlq,bqswl[0],bqswl[3])[0]['quan']
        if cq==0:
            database.execute(sqlu,bqswl[1],bqswl[0],bqswl[3])
        else:
            cq=cq+int(bqswl[1])
            database.execute(sqlu,cq,bqswl[0],bqswl[3])
            database.execute("delete from tobedealwithtable where saon=%s and banu=%s ",saon,bqswl[0])
        return True
    except:
        return False
###############################################################################################################################################
def tobeconfirmdelieverydataframe(saon):
	saon=saon
	sqlm="select * from tobedealwithtable where saon=%s"
	recm=database.query(sqlm,saon)
	drecm=pd.DataFrame(recm)
	drecm=drecm.set_index("id")
	drecm=drecm.reset_index(drop=True)

	return drecm

###############################################################################################################################################
def pallets(saon):
	saon=saon

	sqlt="select * from tobedealwithtable where saon=%s"
	rect=database.query(sqlt,saon)
	sqld="select * from delieverytable where saon=%s"
	recd=database.query(sqld,saon)
	if len(rect)>0:
		recm=rect
	elif len(recd)>0:
		recm=recd
	else:
		recm=''
	if len(recm)>0:
		drecm=pd.DataFrame(recm)
		drecm=drecm.set_index("id")
		drecm=drecm.reset_index(drop=True)
		acount225=drecm[drecm.past==225]["quan"].sum()+drecm[drecm.past==215]["quan"].sum()+drecm[drecm.past==235]["quan"].sum()
		ps=calculatepallet(acount225)
		return ps
	else:
		return 0
	

	
###############################################################################################################################################
def calculatepallet(a):
     x=a/4
     if a%4>0:
             x=x+1
     else:
             x=x
     return x


################################################################################################################################################
def databasetabletodataframe(a):
	sql="select * from "+a
	rec=database.query(sql)
	if len(rec)>0:
		df=pd.DataFrame(rec)
		dfd=df.set_index("id")
		return dfd
	else:
		return ""
##################################################################################################################################################
def databasetabletodataframe2(a,b):
	sql="select * from  "+a+" where cuna like  %s"
	rec=database.query(sql,b)
	if len(rec)>0:
		df=pd.DataFrame(rec)
		dfd=df.set_index("id")
		return dfd
	else:
		return ""

#################################################################################################################################################
def rectodataframe(a):
	df=pd.DataFrame(a)
	dfd=df.set_index("id")
	return dfd
#################################################################################################################################################
####清除字典中的空值
def fillnulldict(d):
	i=0
	while i<len(d):
		a=d.keys()[i]
		b=d[a]
		if type(b)=='NoneType':
			d[a]="无".decode('utf8')
		else:
			pass
		i=i+1
		return d
	

#######用e字典中的值替换d中对应的e字典中的键名
def loopreplaceindocx(d,e):
     i=0
     while i<len(e):
             a=e.keys()[i]
             b=e[a]
             findandreplace(d,a,b)
             i=i+1
 

################################################################################################################################################
###从d文档中用b替换a
def findandreplace(d,a,b):
     for p in d.paragraphs:
             if a in p.text:
                     rp=p.text.replace(a,b)
                     p.text=rp
             else:
                     pass
#########################################################################################
####预扣库存
def prereducestock(d):
	if len(d)>0:
		df=d.reset_index(drop=True)
		i=0
		while i<len(d):
			sqlr='select quan from stocktable where banu=%s and wana=%s'
			sql='update stocktable set quan= %s  where banu = %s and wana = %s'
			rquan=database.query(sqlr,df.ix[i]["banu"],df.ix[i]["wana"])
			equan=rquan[0]["quan"]-df.ix[i]["quan"]
			database.execute(sql,equan,df.ix[i]["banu"],df.ix[i]["wana"])
			i=i+1
		return True
	else:
		return False
			
	

##################################################################################################################################################
#根据可出材料明细表写入出货明细表。输入参数有销售单号，出货明细表及业务员和助理名称
def writeintobedealwithtable(a,b,c,d,p,i):
	saon=a
	sape=b
	mape=c
	pric=p
	insf=i
	df=d
	tday=strtoday()
	l=len(df)
	i=0
	if database:
		sql="insert into tobedealwithtable (saon,sape,mape,mana,banu,past,quan,wana,mada,sapr,exda,pagw,insf) values (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
		while i<l:
			database.execute(sql,saon,sape,mape,df.ix[i]["mana"],df.ix[i]["banu"],df.ix[i]["past"],df.ix[i]["quan"],df.ix[i]["wana"],tday,float(pric),df.ix[i]["exda"],df.ix[i]["pagw"],insf)
			i=i+1

	else:
		self.write("数据库连结不上.")
##################################################################################################################################################
### 已出部分货物明细
def partdelieverydetail(saon):	
	sql="select * from tobedealwithtable where saon=%s order by mana asc,past asc"
	if database:
		pdd=database.query(sql,saon)
		pdddf=pd.DataFrame(pdd)
		pdddf=pdddf.reset_index(drop=True)
		#ldf=pdddf.index.tolist()#提取索引并转换成list
	else:
		self.write("can't connect with databse")
	return pdddf
		
		




##################################################################################################################################################
#根据出货材料规格数量仓库确定可出材料及批号明细列表。自动检查数量是否有效。返回值是一个有效期由近到远的DataFrame
def recstock(mana,past,quan,wana):
	mana=mana
	past=past
	quan=quan
	wana=wana
	manaq="%"+mana+"%"
	sql="select * from stocktable where mana like %s   order by right(exda,4) asc ,left(exda,2) asc"
	rec=database.query(sql,manaq)	
	if len(rec):
		dfda=rectodataframe(rec)

		dfd=dfda[(dfda.wana==wana)&(dfda.past==past)&(dfda.quan>0)]
		#dfw=dfd.sort_values(by='exda'.decode('utf-8'),ascending=True)
		dfw=dfd
		df=dfw.reset_index(drop=True)
		qs=(df.quan * df.past).sum()

		
		if qs<quan:
			return ""
		else:
			aloc=findloc(df["quan"],divmod(quan,past)[0])
			ax=aloc[0]
			dfa=df.head(ax+1)
			dfa=df.head(ax+1)
			dfa=dfa.reset_index(drop=True)
			dfa["quan"][ax]=aloc[1]

			

			return dfa
			
	else: return ""
 #####################################################################################################################################
 #######get customer special materal sell price recentest usede in delieverytable.if no result ,return 0.
def getrecentprice(cuna,mana):
    sql='select * from delieverytable where cuna=%s and mana=%s order by owda desc'
    rec=database.query(sql,cuna,mana)
    list=[]
    if len(rec)>0:
        df=pd.DataFrame(rec)
        insf=df.head(1).insf[0].encode("utf-8")
        sapr=df.head(1).sapr[0].encode("utf-8")
        list.append(insf)
        list.append(sapr)
        return list
    else:
        return list
#######################################################################################################################################
#######get one customer all mana  in recent 
def getcustomerallmana(cuna):
    sql="select mana from delieverytable where cuna=%s"
    rec=database.query(sql,cuna)
    ltep=[]
    if len(rec)>0:
        df=pd.DataFrame(rec)
        dfl=df["mana"].tolist()
        dflt=list(set(dfl))
        return dflt
    else:
        return ltep
#############################################################################################################################
######get one customer all price in recent
def getcustomerallprice(cuna):
    lp=getcustomerallmana(cuna)
    rs=dict()
    if len(lp)>0:
        for mana in lp:
            rs[mana]=getrecentprice(cuna,mana)
    else:
        pass
    return rs
        
        
##################################################################################################################
####get cuna from saon
def getcunafromsaon(saon):
    sql="select cuna from saoncustomertable where saon=%s"
    cuna=database.query(sql,saon)[0].cuna.encode("utf-8")
    return cuna
##################################################################################################################################################
#根据出货材料规格数量仓库确定可出材料及批号明细列表。自动检查数量是否有效。返回值是一个有效期由近到远的DataFrame
def recstockb(banu,past,quan,wana):
    #get banu,past,quan and wana infomation.
	banu=banu
	past=past
	quan=quan
	wana=wana
     #select all banu (in fact will only one banu in each wana),order by expire shirt to long.
	sql="select * from stocktable where banu=%s and wana=%s and past=%s   order by right(exda,4) asc ,left(exda,2) asc"
	rec=database.query(sql,banu,wana,past)	
	if len(rec):
		dfda=rectodataframe(rec)

		dfd=dfda[dfda.quan>0]
		dfw=dfd.sort_values(by='exda'.decode('utf-8'),ascending=True)
		df=dfw.reset_index(drop=True)
		qs=df.quan.sum()

		
		if qs<quan:
			return ""
		else:
			aloc=findloc(df["quan"],quan)
			ax=aloc[0]
			dfa=df.head(ax+1)
			dfa=df.head(ax+1)
			dfa=dfa.reset_index(drop=True)
			dfa["quan"][ax]=aloc[1]
			

			return dfa
			
	else: return ""

#################################################################################################################################################
##########显示系统根据要求选出的待出货明细
def showstockinfo(d):
	tableheader='您确定要出材料的总重量'.decode('utf-8')
	kgamount="公斤。总桶数是:".decode('utf-8')
	materialname="材料名称".decode('utf-8')
	batch="批号".decode('utf-8')
	stand="规格".decode('utf-8')
	packaging="包装".decode('utf-8')
	quantity="数量".decode('utf-8')
	expiredate="有效期".decode('utf-8')
	warehouse="仓库".decode('utf-8')
	df=d		
	head="<center><table border=2px><tr><td>"+materialname+"</td><td>"+batch+"</td><td>"+stand+"</td><td>"+packaging+"</td><td>"+quantity+"</td><td>"+expiredate+"</td><td>"+warehouse+"</td></tr>"
	content=''
	end="</table></center>"
	i=0
	l=len(df)
	if len(df):
		while i<l:
		
			content=content+"<tr><td>"+df.ix[i]["mana"]+"</td><td>"+df.ix[i]["banu"]+"</td><td>"+str(df.ix[i]["past"])+"</td><td>"+df.ix[i]["pana"]+"</td><td>"+str(df.ix[i]["quan"])+"</td><td>"+str(df.ix[i]["exda"])[0:10]+"</td><td>"+df.ix[i]["wana"]+"</td></tr>"
			i=i+1
		quantity=(df.quan*df.past).sum()
		ammount=df.quan.sum()


		return ("<center><h1>"+tableheader+str(quantity)+kgamount+str(ammount)+"</h1></center>"+head+content+end)


	else:return ("<center><h1>您查询的结果为空.</h1></center>")
#########################################################################################
####从筛选结果根据材料数量选择到第几行i为止，最后一行的数量是多少d。
def findloc(a,b):
	
	rec=[]
	i=0
	c=b-a[0]
	if not c>0:
		rec.append(i)
		rec.append(b)
		return rec
	else:
		while c>0:
			c=c-a[i+1]
			i=i+1
		d=a[i]+c
		rec.append(i)
		rec.append(d)
		return rec

#########################################################################################
#####取得当前日期的字符格式
def strtoday():
	return datetime.datetime.now().strftime("%Y%m%d")

#########################################################################################
'''#####生成出货文件号生成结果rl[0]是可能的发货号列表，rl[1]是新的a保持a是最后一个号特征
def updateda(a):
  global da
	if a[0:6]==strtoday()[0:6]:
		ac=str(int(a[6:9])+1)
		l=len(ac)
		if l==1:
			d="00"+ac
		elif l==2:
			d="0"+ac
		elif l==3:
			d=ac
		else:
			pass
	else:
		d="001"
	na=strtoday()[0:6]+d
	nas=na[2:9]
	nh=nas[0:2]
	nt=nas[2:7]
	l=[]
	l.append("G"+nh+"B"+nt)
	l.append("S"+nh+"B"+nt)
	l.append("B"+nh+"B"+nt)
	l.append("X"+nh+"B"+nt)
	l.append("D"+nh+"B"+nt)
	da=na
	return l
#########################################################################################
#####生成进货文件号生成结果rl[0]是可能的发货号列表，rl[1]是新的a保持a是最后一个号特征
def updateia(a):
  global ia
	if a[0:6]==strtoday()[0:6]:
		ac=str(int(a[6:9])+1)
		l=len(ac)
		if l==1:
			d="00"+ac
		elif l==2:
			d="0"+ac
		elif l==3:
			d=ac
		else:
			pass
	else:
		d="001"
	na=strtoday()[0:6]+d
	nas=na[2:9]
	nh=nas[0:2]
	nt=nas[2:7]
	l=[]
	l.append("X"+nh+"B"+nt)
	l.append("S"+nh+"B"+nt)
	l.append("B"+nh+"B"+nt)
	l.append("G"+nh+"B"+nt)
	ia=na
	return l
'''
#########################################################################################
########根据货运地区，毛重，是否危险品以及运费是自付还是代垫计算运费
def calcfreight(a,b,c,d):
        frar=a
	frwe=b



	sfdl=c
	sfcu=d
	dresult={}
	dresult["frar"]=frar	
	dresult["frwe"]=frwe
	#dresult["sfdl"]=sfdl	
	#dresult["sfcu"]=sfcu		
	if sfdl=="gdanger":
		dresult["sfdl"]="危险品运输"	
		if sfcu=="D":
			dresult["sfcu"]="运费代垫"
			dz=1.04
		elif sfcu=="Z":
			dresult["sfcu"]="运费自付"
			dz=1.03
		else:pass

		gdfp=getdfreightpricecolumn(float(frwe))
		sqlp="select "+gdfp+" from dangerfreighttable where frar='"+frar+"'"
		try:
			price=database.query(sqlp)
			dresult["frpr"]=round(round(round(round(price[0][gdfp],2)*dz,2)/1.11,2)*1.1,2)
			dresult["frei"]= round(getdfreightweight(float(frwe))*round(round(round(round(price[0][gdfp],2)*dz,2)/1.11,2)*1.1,2),2)

			return dresult
		except:
			return 0

	elif sfdl=="gnormal":
		dresult["sfdl"]="普货运输"
		
		if sfcu=="D":
			dresult["sfcu"]="运费代垫"
			dz=1.04

		elif sfcu=="Z":
			dresult["sfcu"]="运费自付"
			dz=1.03
		else:pass

		gnfp=getnfreightpricecolumn(float(frwe),sfcu)
		sqlp="select "+gnfp+" from normalfreighttable where frar='"+frar+"'"
		try:
			price=database.query(sqlp)
			#dresult["frpr"]=round(round(price[0][gnfp],2)*dz,2)
            ##new rate 11 to 10
			dresult["frpr"]=round(round(round(round(price[0][gnfp],2)*dz,2)/1.11,2)*1.1,2)
			dresult["frei"]=round(getnfreightweight(float(frwe))*round(round(round(round(price[0][gnfp],2)*dz,2)/1.11,2)*1.1,2),2)
			return dresult
		except:
			return 0

		
	elif sfdl=="gnormalgd":
		dresult["sfdl"]="广州发广东自付普货运输"
		if sfcu=="D":
			return 0
		elif sfcu=="Z":
			
			dresult["sfcu"]="运费自付"
			#dresult["frpr"]=51.5
            ##new rate 11 to 10
			dresult["frpr"]=51.04
			dresult["frei"]= round(getnfreightweight(float(frwe))*51.04,2)
			return dresult
		else:
			pass

	else:
		pass

#########################################################################################
########取得普货运价档位
def getnfreightpricecolumn(a,b):
	if a>=12:
		return (b+"12T")
	elif a<12 and a>=8:
		return (b+"8T")
	elif a<8 and a>=5:
		return (b+"5T")
	elif a<5 and a>=3:
		return (b+"3T")
	elif a<3 and a>=1:
		return (b+"1T")
	elif a<1 and a>=0:
		return (b+"0T")
	else:pass
#########################################################################################
########取得危品运价档位
def getdfreightpricecolumn(a):
	if a>=5:
		return ("DZ510")
	elif a<5 and a>=3:
		return ("DZ35")
	elif a<3 and a>=0:
		return ("DZ13")
	else:pass
#########################################################################################
#####对危品计费毛重处理
def getdfreightweight(a):
	if a<1:
		return 1
	elif a>=1:
		return round(a,1)
	else:pass
#########################################################################################
######对普货计费毛重处理。
def getnfreightweight(a):
	if a<0.2:
		return 0.2
	elif a>=0.2 and a<0.5:
		return round(a,2)
	elif a>=0.5:
		return round(a,1)
	else:pass
#########################################################################################
#####d为DataFrame,a是单个列值string,对单个列的内容去重。返回list
def dropduplicates(d,a):
	result=d[a].drop_duplicates().tolist()
	return result
#########################################################################################
#####取得现有各库存名列表
def eachstockname():
	rec=database.query("select wana from stocktable")
	df=pd.DataFrame(rec)
	rl=df["wana"].drop_duplicates().tolist()
	return rl
#########################################################################################
#####取得现有各库存
def eachstock(stockname):
	rec=database.query("select * from stocktable where wana=%s",stockname)
	df=pd.DataFrame(rec)
	df=df.reset_index(drop=True)
	return df	
#########################################################################################
#####取得现有各库存材料列表
def eachstockml(stockdf):
	dfml=stockdf["mana"].drop_duplicates().tolist()
	return dfml
#########################################################################################
#####取得现有各材料在各库存的汇总DF
def eachmanastock(eachstock,mana):
	df=eachstock[eachstock["mana"]==mana]
	return df
#########################################################################################
#####取得现有各材料在各库存的汇总DF中的包装规格列表
def eachmanastockpast(eachmanastock):
	emspl=eachmanastock["past"].drop_duplicates().tolist()
	return emspl
#########################################################################################
#####取得现有各材料在各库存的各包装规格汇总表:材料名称,库别,包装规格,同规格重量公斤,汇总全库重量,汇总各库重量.
def detailstock():
	rec=database.query("select mana from stocktable")
	df=pd.DataFrame(rec)
	manal=df["mana"].drop_duplicates().tolist()
	rl=[]
	###detail stock list 
	stocklist=eachstockname()
	stock=[]
	for sl in stocklist:
		sd={}
		es=eachstock(sl)
		sd[sl]=es
		stock.append(sd)
		

	##each mana
	for ml in manal:
		dml={}
		dml["mana"]=ml
		##eachmana detail in eachstock
		for estock in stock:
			stockname=estock.keys()[0]
			sndf=estock.values()[0]
			sndfml=sndf[sndf.mana==ml]
			if len(sndfml)>0:
				dml[stockname]=(sndfml.quan*sndfml.past).sum()
				
			else:
				dml[stockname]=0
		rl.append(dml)
	return rl
				

#########################################################################################
####取得待处理出货订单号列表
def getsaonlist(d):
	sl=dropduplicates(d,"saon")
	return sl
	
#########################################################################################
###取得等处理发货(tobedealwithtable)明细表，转换成DataFrame格式.
def getinfotobedealwithtable():
	try:
		rec=database.query("select * from tobedealwithtable")
		d=pd.DataFrame(rec)
		#df=d.set_index("id")
		df=d.drop(["id"],axis=1)
		df=df.reset_index(drop=True)
		return df
	except:
		return "<center><h1>数据库连结失败。请联系管理员。</h1></center>"			

#########################################################################################
###取得等处理发货单个发货号对应的(tobedealwithtable)明细表，转换成DataFrame格式.
def getonetobedealwithtable(s):
	saon=s
	try:
		rec=database.query("select * from tobedealwithtable where saon=%s",saon)
		d=pd.DataFrame(rec)
		#df=d.set_index("id")
		df=d.drop(["id"],axis=1)
		df=df.reset_index(drop=True)
		return df
	except:
		return "<center><h1>数据库连结失败。请联系管理员。</h1></center>"	

#########################################################################################
##根据查询截止年月计算库存符合条件库存
def exdarq(n):
	dy=n[0:4]
	dm=n[4:6]
	sql="select * from stocktable where (left(exda,2)<=%s  and right(exda,4)<=%s ) or (right(exda,4)<%s) order by mana, right(exda,4) asc ,left(exda,2) asc"
	rec=database.query(sql,dm,dy,dy)
	if len(rec)>0:
		return rec
	else:
		return ""
#########################################################################################
###根据有效期限N个月计算查询截止年月
def exdashow(n):
	d=strtoday()[0:6]
	dy=d[0:4]
	dm=d[4:6]
	c=int(dm)+n
	if c<10:
		cs="0"+str(c)
	elif c>9:
		if c<13:
			cs=str(c)
		else:
			cs="0"+str(c-12)
			dy=str(int(dy)+1)
	else:
		pass
	return (dy+cs)
#################################################################################################################
def getfrarall():
	sqln="select frar from normalfreighttable"
	sqld="select frar from dangerfreighttable"
	rsl=[]
	recn=database.query(sqln)
	for recno in recn:
		rsl.append(recno)
	recd=database.query(sqld)
	for recdo in recd:
		rsl.append(recdo)
	return rsl




